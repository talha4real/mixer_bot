import discord
from discord.ext import commands
import gdown
from moviepy.editor import VideoFileClip
from moviepy.video.fx import colorx
import numpy as np
from g_drive_service import GoogleDriveService
from googleapiclient.http import MediaFileUpload
from docx import Document
import random
import re
from moviepy.video.fx import all as vfx_all
from PIL import Image
from io import BytesIO
import streamlink

captions = []
hashtags = []
doc_mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
doc = Document('./video_captions/Video_Captions.doc')


hash_doc = Document('./video_captions/Hashtags.docx')

tiktok_bio = Document('./bios/tiktok.docx')
tiktok_bio_text = '\n\n'.join([paragraph.text for paragraph in tiktok_bio.paragraphs])
tiktok_bio_points = re.split(r'\b\d+\.\s', tiktok_bio_text)[1:]


instagram_bio = Document('./bios/instagram.docx')
instagram_bio_text = '\n\n'.join([paragraph.text for paragraph in instagram_bio.paragraphs])
instagram_bio_points = re.split(r'\b\d+\.\s', instagram_bio_text)[1:]


hash_text = '\n'.join([paragraph.text for paragraph in hash_doc.paragraphs])
hash_lines = hash_text.split('\n')

for line in hash_lines:
    if line.strip():  # Check if the line is not empty after removing whitespace
        hashtags.append(line)





text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
lines = re.split(r'\b\d+\.\s', text)[1:]




for line in lines[1:]:
    if line.strip():  # Check if the line is not empty after removing whitespace
        captions.append(line)



def get_folder_id(service, folder_name):
    results = service.files().list(q=f"name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder'",
                                   fields="files(id)").execute()
    items = results.get('files', [])
    if items:
        return items[0]['id']
    else:
        return None
    
def create_folder(service, folder_name, parent_folder_id=None):
    folder_metadata = {
        'name': folder_name,
        'mimeType': 'application/vnd.google-apps.folder'
    }
    print(folder_metadata)
    if parent_folder_id:
        print(parent_folder_id)
        folder_metadata['parents'] = [parent_folder_id]

    folder = service.files().create(body=folder_metadata,
                                    fields='id').execute()
    return folder.get('id')

def upload_file(service, folder_id, file_name,name, ftype='video/mp4'):
    try:
        media_body = MediaFileUpload(file_name, mimetype=ftype, resumable=True)
        file_metadata = {
            'name': name,
            'parents': [folder_id]
        }
        file = service.files().create(body=file_metadata,
                                    media_body=media_body,
                                    fields='id').execute()
    except Exception as e:
        print("An error occurred while uploading the file:")
    
def add_noise(frame):
    noise = np.random.normal(scale=10, size=frame.shape).astype(np.uint8)
    noisy_frame = np.clip(frame + noise, 0, 255)
    return noisy_frame

intents = discord.Intents.all()

bot = commands.Bot(command_prefix='!', intents=intents)

# Dictionary to store user submitted links
user_links = {}

def apply_random_effects(image):
    # Apply your effects to the image here
    modified_image = image.copy()

    # Apply colorx effect
    random_factor_colorx = random.uniform(0.5, 1.5)
    modified_image = apply_colorx_effect(modified_image, random_factor_colorx)

    return modified_image

def apply_colorx_effect(image, factor):
    # Convert the image to a NumPy array
    np_image = np.array(image)

    # Apply color adjustment
    modified_np_image = np.minimum(255, (factor * np_image)).astype('uint8')

    # Convert the NumPy array back to an image
    modified_image = Image.fromarray(modified_np_image)

    return modified_image


def image_to_bytes(image):
    # Convert the image to bytes
    buffered = BytesIO()
    image.save(buffered, format='PNG')
    buffered.seek(0)
    return buffered


@bot.command()
async def mix(ctx, platform: str):
    if platform.lower() == 'tiktok' or platform.lower() == 'insta':
        embed = discord.Embed(
            title="Upload Videos for TikTok Mix",
            description="Please upload vidoes on this link: [Streamable](https://streamable.com/). Once Uploaded, copy the video urls and use the command `!begin [video_link] [video_link]` to start mixing.",
            color=discord.Color.blue()
        )
        await ctx.send(embed=embed)
    elif platform.lower()=="pfp":
        if ctx.message.attachments:
            # Get the first attached image
            attachment = ctx.message.attachments[0]

            # Download the image
            response = await attachment.read()
            image = Image.open(BytesIO(response))

            # Apply effects
            modified_image = apply_random_effects(image)

            # Convert the modified image to bytes
            modified_image_bytes = image_to_bytes(modified_image)

            # Send the modified image back
            await ctx.send(file=discord.File(modified_image_bytes, filename='modified_image.png'))
        else:
            await ctx.send("Please send an image alongside the command")


    else:
        await ctx.send("Invalid platform. Use `!mix tiktok`.")



def get_files(folder_name):
    selected_fields="files(id,name,webViewLink)"
    g_drive_service=GoogleDriveService().build()
    list_file=g_drive_service.files().list(fields=selected_fields).execute()
    folder_id = get_folder_id(g_drive_service, folder_name)
    if folder_id:
        query = f"'{folder_id}' in parents"
        results = g_drive_service.files().list(q=query).execute()
        files = results.get('files', [])

        if files:
            return {"files": files}
        else:
            return {"error": "No files found in the specified folder."}
    else:
        return {"error": f"Folder '{folder_name}' not found in Google Drive."}


parent_folder_id = "1N4i1DnEQeR_vCDLrpW9ZlFWW7FFUPO4s"

def mix_video(video,folder_id):

    video_clip = VideoFileClip(video)
        
    modified_clip = video_clip.fl_image(add_noise)


    random_factor = random.uniform(0.5, 1.5)
    if random.random() < 1/3:
        modified_clip = video_clip.fx(colorx.colorx, random_factor)

    modified_clip = video_clip.fx(vfx_all.lum_contrast, random_factor)

    modified_clip = vfx_all.colorx(video_clip, random_factor)


    if random.random() < 0.2:
        modified_clip = vfx_all.blackwhite(video_clip)

    fadein_duration = random.uniform(1, 5)

    fadein_clip = vfx_all.fadein(video_clip, fadein_duration)

    modified_clip = fadein_clip.fx(vfx_all.colorx, random_factor)

    # Output video path
    output_path = f"./output/{video}"

    # Write the modified clip to an output file
    modified_clip.write_videofile(output_path, codec="libx264", audio_codec="aac")
 
    # Close the video clips
    video_clip.close()
    modified_clip.close()


  

g_drive_service=GoogleDriveService().build()

@bot.command()
async def submit(ctx, folder_name: str):
    files = get_files(folder_name)
    if( len(files["files"])>0):
        loading_message = await ctx.send("Mixing... ⏳")
        doc = Document()
        folder_id = create_folder(g_drive_service, f"result_{folder_name}",parent_folder_id)
        for video in files["files"]:
            if video["mimeType"] == "video/mp4":
                mix_video(video,folder_id)
                upload_file(g_drive_service,folder_id, f"./output/{video['name']}",video["name"])
                random_caption = random.choice(captions)
                # Remove the chosen caption from the array
                captions.remove(random_caption)
                doc.add_paragraph(f"Video Name: {video['name']}")
                doc.add_paragraph(random_caption)
                random_hashtags = random.sample(hashtags, 5)
                hashtags_text = ' '.join(random_hashtags)

                doc.add_paragraph(hashtags_text)



        embed = discord.Embed(
            title="Mix Completed",
            description=f"You can find the edited videos in the drive below: \n [Drive](https://drive.google.com/drive/folders/{folder_id}).",
            color=discord.Color.blue()
        )


        doc.save('./output/captions.docx')
        upload_file(g_drive_service,folder_id,'./output/captions.docx',"captions.docx",doc_mime_type)
        await loading_message.edit(content="Mix has been completed! ✅")

        await ctx.send(embed=embed)
        # await ctx.send(f"https://drive.google.com/drive/folders/{folder_id}")
    else:
        await ctx.send("No Videos Found")


@bot.command()
async def backup(ctx, folder_name: str):
    files = get_files(folder_name)
    print(files)
    if( len(files["files"])>0):
        loading_message = await ctx.send("Mixing... ⏳")
        doc = Document()
        folder_id = create_folder(g_drive_service, f"result_{folder_name}",parent_folder_id)
        for video in files["files"]:
            if video["mimeType"] == "video/mp4":
                mix_video(video,folder_id)
                upload_file(g_drive_service,folder_id, f"./output/{video['name']}",video["name"])
                random_caption = random.choice(captions)
                # Remove the chosen caption from the array
                captions.remove(random_caption)
                doc.add_paragraph(f"Video Name: {video['name']}")
                doc.add_paragraph(random_caption)
                random_hashtags = random.sample(hashtags, 5)
                hashtags_text = ' '.join(random_hashtags)

                doc.add_paragraph(hashtags_text)



        embed = discord.Embed(
            title="Mix Completed",
            description=f"You can find the edited videos in the drive below: \n [Drive](https://drive.google.com/drive/folders/{folder_id}).",
            color=discord.Color.blue()
        )


        doc.save('./output/captions.docx')
        upload_file(g_drive_service,folder_id,'./output/captions.docx',"captions.docx",doc_mime_type)
        await loading_message.edit(content="Mix has been completed! ✅")

        await ctx.send(embed=embed)
        # await ctx.send(f"https://drive.google.com/drive/folders/{folder_id}")
    else:
        await ctx.send("No Videos Found")

@bot.command()
async def begin(ctx, *links):
    if len(links)>0:
        loading_message = await ctx.send("Mixing... ⏳")
        doc = Document()
        folder_id = create_folder(g_drive_service, f"result_mixer",parent_folder_id)
        for link in links:
            session = streamlink.Streamlink()
            streams = session.streams(link)
            if not streams:
                print("No streams found for the provided URL.")
            else:
                # Choose the best stream (highest quality)
                stream = streams["best"]
                
                # Get the video URL
                video_url = stream.url
                
                # Download the video
                response = session.http.get(video_url, stream=True)
                video_identifier = link.split("/")[-1]
                video_filename = f"{video_identifier}.mp4"
                with open(video_filename, "wb") as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                        
                print("Video downloaded as", video_filename)
                mix_video(video_filename,folder_id)
                upload_file(g_drive_service,folder_id, f"./output/{video_filename}",video_filename)
                random_caption = random.choice(captions)
                # Remove the chosen caption from the array
                captions.remove(random_caption)
                doc.add_paragraph(f"Video Name: {video_filename}")
                doc.add_paragraph(random_caption)
                random_hashtags = random.sample(hashtags, 5)
                hashtags_text = ' '.join(random_hashtags)

                doc.add_paragraph(hashtags_text)
            
        embed = discord.Embed(
            title="Mix Completed",
            description=f"You can find the edited videos in the drive below: \n [Drive](https://drive.google.com/drive/folders/{folder_id}).",
            color=discord.Color.blue()
        )

        doc.save('./output/captions.docx')
        upload_file(g_drive_service,folder_id,'./output/captions.docx',"captions.docx",doc_mime_type)
        await loading_message.edit(content="Mix has been completed! ✅")

        await ctx.send(embed=embed)
        # await ctx.send(f"https://drive.google.com/drive/folders/{folder_id}")
    else:
        await ctx.send("No Videos Found")
  

@bot.command()
async def caption(ctx):
    random_caption_1 = random.choice(captions)
    captions.remove(random_caption_1)
    random_caption_2 = random.choice(captions)
    captions.remove(random_caption_2)

    random_caption_3 = random.choice(captions)
    captions.remove(random_caption_3)


    embed = discord.Embed(
            title="Captions",
            description="Random Captions",
            color=discord.Color.blue()
        )
    embed.add_field(name="Caption 1", value=random_caption_1, inline=False)
    embed.add_field(name="Caption 2", value=random_caption_2, inline=False)
    embed.add_field(name="Caption 3", value=random_caption_3, inline=False)

    await ctx.send(embed=embed)     


# @bot.command()
# async def insta(ctx, platform: str):
#     if platform.lower() == 'bio':
#         bio_point =  random.choice(instagram_bio_points).replace("\n", "") 

#         embed = discord.Embed(
#             title="Instagram Bio",
#             description=bio_point,
#             color=discord.Color.blue()
#         )
#         await ctx.send(embed=embed)
    

@bot.command()
async def bio(ctx, platform: str):
    if platform.lower() == 'tiktok':

        bio_point =  random.choice(tiktok_bio_points).replace("\n", "") 

        embed = discord.Embed(
            title="TikTok Bio",
            description=bio_point,
            color=discord.Color.blue()
        )
        await ctx.send(embed=embed)
    elif platform.lower() == 'instagram':

        bio_point =  random.choice(instagram_bio_points).replace("\n", "") 

        embed = discord.Embed(
            title="Instagram Bio",
            description=bio_point,
            color=discord.Color.blue()
        )
        await ctx.send(embed=embed)

    else:
        await ctx.send("Invalid platform. Use `!mix tiktok`.")


@bot.command()
async def commands(ctx):
    # Create an embed for the help message
    embed = discord.Embed(title="Bot Help", description="List of available commands:", color=discord.Color.blue())

    # Add fields for each command with their descriptions
    embed.add_field(name="!mix tiktok",
                    value="Run the Command To see the steps.",
                    inline=False)
    
   # Add fields for each command with their descriptions
    embed.add_field(name="!mix insta",
                    value="Run the Command To see the steps.",
                    inline=False)
    
    embed.add_field(name="!mix pfp",
                    value="Person uploads image to the bot directly. Takes it -> changes meta data and returns the edited image.",
                    inline=False)

    embed.add_field(name="!bio tiktok",
                    value="Function: Returns a caption for the bio of an account for TikTok.",
                    inline=False)

    embed.add_field(name="!bio instagram",
                    value="Function: Returns a caption for the bio of an account for Instagram.",
                    inline=False)

    embed.add_field(name="!caption",
                    value="Returns 3 captions from a pool. These captions will be used for the post.",
                    inline=False)

 

    # Send the embed as a message
    await ctx.send(embed=embed)

@bot.event
async def on_ready():
    print(f'Logged in as {bot.user.name}')




bot.run('MTE0NDQ2OTgzNDM0ODkwNDUzMQ.GolKpW.5TeN1c1lH3JE3Lw2syzBmFbjy1F8WKhACusU20')
